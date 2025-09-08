# modules/file_processing.py - UPDATED for integrated chat flow
# Fixes Entry #4: Now keeps file analysis in the conversational thread
# FIXED: Task extraction with mandatory format enforcement

import os
import io
import tempfile
from PIL import Image
import fitz
import docx
import markdown
from markupsafe import Markup
from flask import request, redirect, jsonify, current_app
from utils.ghostline_engine import generate_response
from utils.rag_basic import is_ready
from modules.database import save_conversation_enhanced, track_uploaded_file
from modules.brain import enhanced_retrieve

PROJECTS = [
    'Personal Operating Manual','AMCF','BCDodgeme','Rose and Angel','Meals N Feelz',
    'TV Signals','Damn It Carl','HalalBot','Kitchen','Health','Side Quests'
]

CHAT_MODEL = os.getenv("CHAT_MODEL", os.getenv("OPENROUTER_MODEL", "openrouter/auto"))

def setup_easyocr_environment():
    """Setup EasyOCR environment variables"""
    import os
    os.environ['EASYOCR_MODULE_PATH'] = '/tmp/easyocr_models'
    print(f"EasyOCR model path set to: {os.environ.get('EASYOCR_MODULE_PATH', 'default')}")

def markdown_filter(text):
    """Convert markdown to HTML for template rendering"""
    if not text:
        return ""
    html = markdown.markdown(text)
    return Markup(html)

def get_project_specific_analysis_context(project):
    """Return project-specific context for better document analysis"""
    project_contexts = {
        'Personal Operating Manual': "personal productivity, self-improvement, habits, and life optimization",
        'AMCF': "business operations, client work, project management, and deliverables",
        'BCDodgeme': "game development, mechanics, user experience, and technical implementation",
        'Rose and Angel': "relationship management, personal connections, and social dynamics",
        'Meals N Feelz': "nutrition, cooking, meal planning, and wellness tracking",
        'TV Signals': "media projects, content creation, and entertainment development",
        'Damn It Carl': "personal challenges, problem-solving, and self-accountability",
        'HalalBot': "technical projects, automation, and system development",
        'Kitchen': "cooking experiments, recipe development, and culinary exploration",
        'Health': "medical tracking, fitness goals, and wellness monitoring",
        'Side Quests': "experimental projects, learning goals, and creative pursuits"
    }
    
    return project_contexts.get(project, "general productivity and organization")

def select_optimal_voices_for_document(filename, content_preview, project):
    """Select the best voice(s) based on document characteristics"""
    
    filename_lower = filename.lower()
    content_lower = content_preview.lower()
    
    # Task-heavy documents - SyntaxPrime for organization
    if any(keyword in content_lower for keyword in [
        'todo', 'task', 'action', 'deadline', 'meeting', 'project',
        'deliverable', 'milestone', 'schedule', 'agenda'
    ]):
        return ['SyntaxPrime']
    
    # Creative or experimental content - SyntaxBot for perspective
    if any(keyword in content_lower for keyword in [
        'creative', 'idea', 'concept', 'design', 'art', 'story',
        'experiment', 'prototype', 'brainstorm'
    ]) or project in ['Side Quests', 'TV Signals', 'BCDodgeme']:
        return ['SyntaxBot']
    
    # Technical/logical content - Nil.exe for analysis
    if any(keyword in content_lower for keyword in [
        'code', 'technical', 'specification', 'requirements', 'bug',
        'system', 'architecture', 'algorithm', 'data'
    ]) or filename_lower.endswith(('.py', '.js', '.json', '.xml', '.csv')):
        return ['Nil.exe']
    
    # Personal/relationship content - GGPT for warmth
    if any(keyword in content_lower for keyword in [
        'personal', 'relationship', 'family', 'friend', 'emotion',
        'feeling', 'health', 'wellness', 'habit'
    ]) or project in ['Rose and Angel', 'Health', 'Meals N Feelz']:
        return ['GGPT']
    
    # Default to SyntaxPrime for general analysis
    return ['SyntaxPrime']

def generate_contextual_analysis_prompt(file, text, project, vision_description=None):
    """Generate analysis prompt that's aware of project context and user needs"""
    
    project_context = get_project_specific_analysis_context(project)
    filename = file.filename
    
    if vision_description:
        # Image with vision analysis
        return f"""I uploaded '{filename}' to my {project} project. 

The OCR didn't pick up much text, but here's what I can see in the image:
{vision_description}

Given that this is for {project_context}, what should I do with this? Any insights that relate to what I'm working on in this area?"""

    elif filename.endswith(('.pdf', '.docx')):
        # Document analysis with project awareness
        return f"""I uploaded '{filename}' to my {project} project.

CONTENT:
{text}

This is related to {project_context}. What do I need to know from this document? 

Specifically:
- Any actions I should take?
- Information that affects my current work in this area?
- Deadlines or commitments I need to track?
- Key insights that matter for {project}?

Skip the summary unless it's actually useful - just tell me what I should do with this information."""

    else:
        # General file analysis
        return f"""I uploaded '{filename}' to {project}.

CONTENT:
{text}

This is going into my {project_context} tracking. What's worth paying attention to here? Any actions needed or just reference material?

Given the context of {project}, does this change anything about what I'm working on or planning?"""

def generate_document_analysis_response(analysis_prompt, project, voices, random_toggle):
    """Generate personality-driven document analysis responses"""
    
    # Enhance the prompt to encourage natural conversation
    enhanced_prompt = f"""Carl just uploaded a document for analysis. Respond naturally as his AI assistant - don't be overly formal or academic.

{analysis_prompt}

Remember:
- Stay in your character voice/personality
- Be practical and actionable, not theoretical  
- If Carl needs to do something, tell him clearly
- If it's just reference material, say so
- Use your normal conversational tone
- Don't force formal structures unless the content actually demands them
- Focus on what Carl actually needs to know, not showing off analysis skills"""

    # Use existing response generation with enhanced context
    try:
        retrieval_ctx = enhanced_retrieve(analysis_prompt, k=3, project=project) if is_ready() else []
        
        # Generate response using existing personality system
        response_data = generate_response(
            enhanced_prompt, voices, random_toggle,
            project=project, model=CHAT_MODEL, retrieval_context=retrieval_ctx
        )
        
        return response_data
        
    except Exception as e:
        print(f"Enhanced analysis failed, falling back to basic: {e}")
        # Fallback to original generation
        return generate_response(
            analysis_prompt, voices, random_toggle,
            project=project, model=CHAT_MODEL, retrieval_context=[]
        )

def process_image_ocr(file_stream, filename):
    """Process image with EasyOCR, handling model download issues"""
    try:
        import easyocr
        import numpy as np
        
        print(f"Starting OCR processing for: {filename}")
        
        # Reset stream position
        file_stream.seek(0)
        
        # Open and convert image
        img = Image.open(file_stream)
        
        # Convert to RGB if necessary
        if img.mode != 'RGB':
            img = img.convert('RGB')
        
        img_array = np.array(img)
        print(f"Image loaded: {img.size}, mode: {img.mode}")
        
        # Initialize EasyOCR with error handling for model downloads
        try:
            # Try to create reader with custom model path
            reader = easyocr.Reader(
                ['en'],
                gpu=False,  # Disable GPU for server compatibility
                download_enabled=True,  # Allow model downloads
                model_storage_directory=os.environ.get('EASYOCR_MODULE_PATH')
            )
            print("EasyOCR reader initialized successfully")
            
        except Exception as model_error:
            print(f"EasyOCR model initialization failed: {model_error}")
            
            # Fallback: try without custom directory
            reader = easyocr.Reader(['en'], gpu=False)
            print("EasyOCR reader initialized with default settings")
        
        # Perform OCR
        results = reader.readtext(img_array)
        
        if results:
            text = '\n'.join([result[1] for result in results if result[1].strip()])
            print(f"OCR extracted {len(results)} text regions, {len(text)} characters")
            return text
        else:
            print("No OCR results found")
            return "No text detected in image"
            
    except ImportError as e:
        print(f"EasyOCR not installed: {e}")
        raise Exception("EasyOCR not installed. Please install with: pip install easyocr opencv-python-headless")
    
    except Exception as e:
        print(f"OCR processing failed: {e}")
        raise Exception(f"OCR processing failed: {str(e)}")

def analyze_image_with_vision(file_stream, filename):
    """Analyze image using GPT-4 Vision when OCR results are poor"""
    try:
        import base64
        import requests
        
        # Get OpenAI API key from environment
        openai_api_key = os.getenv('OPENAI_API_KEY')
        if not openai_api_key:
            return "OpenAI API key not configured for vision analysis"
        
        # Reset stream and encode image
        file_stream.seek(0)
        image_data = base64.b64encode(file_stream.read()).decode('utf-8')
        
        # Determine image format for data URL
        file_extension = filename.split('.')[-1].lower()
        mime_type = f"image/{file_extension}" if file_extension in ['png', 'jpg', 'jpeg', 'gif', 'bmp'] else "image/jpeg"
        
        headers = {
            "Authorization": f"Bearer {openai_api_key}",
            "Content-Type": "application/json"
        }
        
        # Create vision analysis prompt
        payload = {
            "model": "gpt-4o",
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": "Analyze this image in detail. If it contains charts, graphs, screenshots, or data visualizations, describe the key insights, trends, and important information. If it's a photo, describe what you see. Be specific and actionable in your analysis."
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:{mime_type};base64,{image_data}",
                                "detail": "high"
                            }
                        }
                    ]
                }
            ],
            "max_tokens": 500,
            "temperature": 0.1
        }
        
        print(f"Sending image to GPT-4 Vision for analysis: {filename}")
        
        response = requests.post(
            "https://api.openai.com/v1/chat/completions",
            headers=headers,
            json=payload,
            timeout=30
        )
        
        if response.status_code == 200:
            result = response.json()
            vision_analysis = result['choices'][0]['message']['content']
            print(f"GPT-4 Vision analysis successful: {len(vision_analysis)} characters")
            return vision_analysis
        else:
            print(f"GPT-4 Vision API error: {response.status_code} - {response.text}")
            return f"Vision analysis failed: API error {response.status_code}"
            
    except Exception as e:
        print(f"Vision analysis failed: {e}")
        return f"Vision analysis error: {str(e)}"

def process_single_file(file, project):
    """Process a single file and return analysis data - IMPROVED VERSION"""
    filename = file.filename.lower()
    text = ""
    
    print(f"Processing file: {filename} for project: {project}")

    # Process different file types
    if filename.endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
        text = process_image_ocr(file.stream, filename)
            
    elif filename.endswith('.pdf'):
        try:
            file.stream.seek(0)
            data = file.read()
            
            if not data:
                raise Exception("PDF file appears to be empty")
            
            doc = fitz.open(stream=data, filetype="pdf")
            
            if doc.page_count == 0:
                raise Exception("PDF has no pages")
            
            text_parts = []
            for page_num in range(doc.page_count):
                page = doc[page_num]
                page_text = page.get_text()
                if page_text.strip():
                    text_parts.append(f"=== Page {page_num + 1} ===\n{page_text}")
            
            text = "\n\n".join(text_parts)
            doc.close()
            
            if not text.strip():
                text = "No text found in PDF (may be image-based or encrypted)"
                
        except Exception as e:
            print(f"PDF processing failed: {e}")
            raise Exception(f"PDF Error: {str(e)}")
            
    elif filename.endswith('.docx'):
        try:
            file.stream.seek(0)
            file_data = file.read()
            
            if not file_data:
                raise Exception("Word document appears to be empty")
            
            file_stream = io.BytesIO(file_data)
            document = docx.Document(file_stream)
            
            paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
            
            tables_text = []
            for table in document.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        tables_text.append(" | ".join(row_text))
            
            all_text = []
            if paragraphs:
                all_text.extend(paragraphs)
            if tables_text:
                all_text.append("\n=== Tables ===")
                all_text.extend(tables_text)
            
            text = "\n".join(all_text)
            
            if not text.strip():
                text = "No readable text found in Word document"
                
        except Exception as e:
            print(f"Word document processing failed: {e}")
            raise Exception(f"Word Document Error: {str(e)}")
            
    else:
        raise Exception(f"Unsupported file type: {filename}. Supported: PNG, JPG, JPEG, GIF, BMP, PDF, DOCX")

    # Truncate if too long
    if len(text) > 15000:
        text = text[:15000] + "\n\n[...Content truncated...]"
    
    # Check if OCR results are meaningful (fallback logic)
    text_words = len(text.split()) if text else 0
    is_image_file = filename.endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp'))
    
    # Generate analysis prompt with enhanced contextual approach
    if is_image_file and text_words < 5:
        # Poor OCR results - switch to GPT-4 Vision analysis
        print(f"OCR extracted only {text_words} words, switching to vision analysis")
        vision_description = analyze_image_with_vision(file.stream, file.filename)
        analysis_prompt = generate_contextual_analysis_prompt(file, text, project, vision_description)
    else:
        # Good text extraction or document files
        analysis_prompt = generate_contextual_analysis_prompt(file, text, project)

    print(f"File processing successful: {len(text)} characters extracted")
    
    return {
        'filename': file.filename,
        'analysis_prompt': analysis_prompt,
        'extracted_text': text,
        'file_size': len(file.read()) if hasattr(file, 'read') else 0
    }

def handle_file_upload():
    """UPDATED: Handle file upload processing with enhanced response generation"""
    try:
        files = request.files.getlist('file')
        if not files or not files[0].filename:
            return jsonify({'error': 'No files uploaded'}), 400
        
        project = request.form.get('project', PROJECTS[0])
        
        # Process all uploaded files
        file_analyses = []
        for file in files:
            if file and file.filename:
                try:
                    file_data = process_single_file(file, project)
                    file_analyses.append(file_data)
                except Exception as e:
                    print(f"Failed to process {file.filename}: {e}")
                    file_analyses.append({
                        'filename': file.filename,
                        'error': str(e)
                    })
        
        if not file_analyses:
            return jsonify({'error': 'No files could be processed'}), 400
        
        # Generate enhanced analysis
        if len(file_analyses) == 1:
            # Single file
            file_data = file_analyses[0]
            if 'error' in file_data:
                return jsonify({'error': f"File processing failed: {file_data['error']}"}), 400
            
            # Smart voice selection
            optimal_voices = select_optimal_voices_for_document(
                file_data['filename'],
                file_data['extracted_text'][:500],
                project
            )
            
            # Generate with improved prompting
            response_data = generate_document_analysis_response(
                file_data['analysis_prompt'],
                project,
                optimal_voices,
                False  # Keep consistent for document analysis
            )
            
            user_message = f"📎 {file_data['filename']}"
            
        else:
            # Multiple files - create natural combined prompt
            successful_files = [f for f in file_analyses if 'error' not in f]
            failed_files = [f for f in file_analyses if 'error' in f]
            
            if not successful_files:
                errors = [f"{f['filename']}: {f['error']}" for f in failed_files]
                return jsonify({'error': f"All files failed: {'; '.join(errors)}"}), 400
            
            # Natural multi-file prompt
            combined_prompt = f"I uploaded {len(successful_files)} files to {project}. Help me understand what I'm looking at and what I should do with them:\n\n"
            
            for i, file_data in enumerate(successful_files, 1):
                combined_prompt += f"FILE {i}: {file_data['filename']}\n"
                combined_prompt += f"{file_data['extracted_text'][:800]}\n\n"
            
            if failed_files:
                combined_prompt += f"Note: {len(failed_files)} files couldn't be processed.\n\n"
            
            combined_prompt += f"Given this is for {project}, what are the key things I should know? Any actions needed across these files?"
            
            response_data = generate_document_analysis_response(
                combined_prompt,
                project,
                ['SyntaxPrime'],  # Use organizing voice for multiple files
                False
            )
            
            filenames = [f['filename'] for f in successful_files]
            user_message = f"📎 {', '.join(filenames)}"
        
        # Save conversation and track files
        save_conversation_enhanced(project, user_message, response_data)
        
        for file_data in file_analyses:
            if 'error' not in file_data:
                filename = file_data['filename']
                file_extension = filename.split('.')[-1].upper() if '.' in filename else 'UNKNOWN'
                content_summary = file_data['extracted_text'][:500] if file_data['extracted_text'] else "No text extracted"
                track_uploaded_file(filename, file_extension, project, content_summary)
        
        return jsonify({
            'success': True,
            'message': 'Files processed successfully',
            'project': project,
            'files_processed': len(successful_files) if len(file_analyses) > 1 else 1
        })
        
    except Exception as e:
        print(f"Upload route failed: {e}")
        current_app.logger.error(f"File upload error: {e}")
        return jsonify({'error': f'Upload failed: {str(e)}'}), 500

# BACKWARD COMPATIBILITY: Keep original function as fallback
def handle_file_upload_legacy():
    """Original file upload handler for backward compatibility"""
    try:
        file = request.files.get('file')
        if not file or not file.filename:
            return "No file uploaded", 400
        
        project = request.form.get('project', PROJECTS[0])
        
        # Process single file using new function
        file_data = process_single_file(file, project)
        
        # Generate AI analysis
        use_voices = ['SyntaxPrime']
        random_toggle = False
        
        try:
            retrieval_ctx = enhanced_retrieve(file_data['analysis_prompt'], k=5) if is_ready() else []
            response_data = generate_response(
                file_data['analysis_prompt'], use_voices, random_toggle,
                project=project, model=CHAT_MODEL, retrieval_context=retrieval_ctx
            )
        except Exception as e:
            print(f"AI analysis failed: {e}")
            response_data = {"SyntaxPrime": f"File processed successfully, but AI analysis failed: {e}"}

        # Save to database
        user_message = f"[File Upload] {file.filename}"
        save_conversation_enhanced(project, user_message, response_data)
        
        # Track the uploaded file
        file_extension = file.filename.split('.')[-1].upper() if '.' in file.filename else 'UNKNOWN'
        content_summary = file_data['extracted_text'][:500] if file_data['extracted_text'] else "No text extracted"
        track_uploaded_file(file.filename, file_extension, project, content_summary)

        # Redirect back to main chat with the analysis
        return redirect(f'/?project={project}#bottom-anchor')
        
    except Exception as e:
        print(f"Legacy upload route failed: {e}")
        return f"Upload Error: {str(e)}", 500
